import re
from io import BytesIO
from typing import List, Dict, Any, Optional

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml import OxmlElement
import logging

__all__ = ["format_document"]

# Настройка логирования
logger = logging.getLogger(__name__)

# Список заголовков первого уровня, которые не нумеруются
NO_NUMERATION_HEADINGS = {
    "РЕФЕРАТ",
    "СОДЕРЖАНИЕ",
    "ТЕРМИНЫ И ОПРЕДЕЛЕНИЯ",
    "ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ",
    "СПИСОК СОКРАЩЕНИЙ",
    "ВВЕДЕНИЕ",
    "ЗАКЛЮЧЕНИЕ",
    "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
    "ПРИЛОЖЕНИЯ",
    "ПРИЛОЖЕНИЕ"
}

def _replace_quotes(text: str) -> str:
    """
    Умная замена кавычек с учётом их позиции в тексте.
    
    Args:
        text (str): Текст для обработки.
    
    Returns:
        str: Текст с заменёнными кавычками.
    """
    quote_pattern = re.compile(r'''(['"“”‘’„‟«»])''')
    
    result = []
    in_quotes = False
    prev_char = ''
    
    for char in text:
        if quote_pattern.match(char):
            if not in_quotes:
                result.append('«')
                in_quotes = True
            else:
                if prev_char in ' \t\n.,;:!?':
                    result.append('«')
                else:
                    result.append('»')
                    in_quotes = False
        else:
            result.append(char)
        prev_char = char
    
    if in_quotes:
        result.append('»')
    
    return ''.join(result)

def _norm_cap(kind: str, cap: str) -> str:
    """Нормализует подпись, удаляя префикс (например, 'Таблица 1 – ' или 'Рисунок 1 – ')."""
    cap_re = {
        "figure": re.compile(
            r'^\s*(?:рисунок|рис\.|figure|fig\.)\s*(?:\d+(?:\.\d+)*)?\s*[-–:]\s*(.+?)(?=\n|$)',
            re.I,
        ),
        "table": re.compile(
            r'^\s*(?:таблица|табл|table|tabl)\s+\d+(?:\.\d+)*\s*[–:]\s*(.+?)(?=\n|$)',
            re.I,
        ),
    }
    match = cap_re[kind].match(cap)
    return match.group(1).strip() if match else cap.strip()

def _is_caption_text(text: str, kind: str) -> bool:
    """Проверяет, соответствует ли текст шаблону подписи для указанного типа (figure или table)."""
    cap_re = {
        "figure": re.compile(
            r'^\s*(?:рисунок|рис\.|figure|fig\.)\s*(?:\d+(?:\.\d+)*)?\s*[-–:]\s*.+?(?=\n|$)',
            re.I,
        ),
        "table": re.compile(
            r'^\s*(?:таблица|табл|table|tabl)\s+\d+(?:\.\d+)*\s*[–:]\s*.+?(?=\n|$)',
            re.I,
        ),
    }
    return bool(cap_re[kind].match(text))

def _strip_heading_number(text: str, level: int) -> str:
    """
    Удаляет старую нумерацию из текста заголовка в зависимости от уровня.

    Args:
        text (str): Текст заголовка.
        level (int): Уровень заголовка (1, 2, 3, 4).

    Returns:
        str: Текст заголовка без нумерации.
    """
    logger.debug(f"Исходный текст заголовка (уровень {level}): '{text}'")

    if level == 1:
        pattern = r'^(?:(?:\d+\s*\.\s*)?(?:ГЛАВА|РАЗДЕЛ|ЧАСТЬ)\s+[0-9IVXLCDM]+\s*\.?\s*|\d+\s*\.\s*\d+\s*\.?\s*|[0-9]+\s*\.?\s*)(.+)$'
        match = re.match(pattern, text, re.IGNORECASE)
        if match:
            cleaned_text = match.group(1).strip()
            cleaned_text = re.sub(r'^\.+|\s+', ' ', cleaned_text).strip()
            logger.debug(f"После удаления нумерации (уровень 1): '{cleaned_text}'")
            return cleaned_text
        logger.debug(f"Заголовок уровня 1 не соответствует шаблону, оставляем как есть: '{text}'")
        return text.strip()
    
    elif level == 2:
        pattern = r'^\d+\.\d+\s*\.?\s*(.+)$'
        match = re.match(pattern, text, re.IGNORECASE)
        if match:
            cleaned_text = match.group(1).strip()
            cleaned_text = re.sub(r'^\.+|\s+', ' ', cleaned_text).strip()
            logger.debug(f"После удаления нумерации (уровень 2): '{cleaned_text}'")
            return cleaned_text
        logger.debug(f"Заголовок уровня 2 не соответствует шаблону, оставляем как есть: '{text}'")
        return text.strip()
    
    elif level == 3:
        pattern = r'^\d+\.\d+\.\d+\s*\.?\s*(.+)$'
        match = re.match(pattern, text, re.IGNORECASE)
        if match:
            cleaned_text = match.group(1).strip()
            cleaned_text = re.sub(r'^\.+|\s+', ' ', cleaned_text).strip()
            logger.debug(f"После удаления нумерации (уровень 3): '{cleaned_text}'")
            return cleaned_text
        logger.debug(f"Заголовок уровня 3 не соответствует шаблону, оставляем как есть: '{text}'")
        return text.strip()
    
    elif level == 4:
        pattern = r'^(\d+\.\d+\.\d+\.\d+|\([A-Za-zА-Яа-яЁё]\)|[A-Za-zА-Яа-яЁё]\.)\s*\.?\s*(.+)$'
        match = re.match(pattern, text, re.IGNORECASE)
        if match:
            cleaned_text = match.group(2).strip()
            cleaned_text = re.sub(r'^\.+|\s+', ' ', cleaned_text).strip()
            logger.debug(f"После удаления нумерации (уровень 4): '{cleaned_text}'")
            return cleaned_text
        logger.debug(f"Заголовок уровня 4 не соответствует шаблону, оставляем как есть: '{text}'")
        return text.strip()
    
    logger.debug(f"Заголовок не обработан (уровень {level}), Ascensor: '{text}'")
    return text.strip()

def _strip_bullet_marker(text: str) -> str:
    """
    Удаляет маркеры (например, -, •, →) из начала текста маркированного списка.

    Args:
        text (str): Текст элемента списка.

    Returns:
        str: Текст без маркера в начале.
    """
    bullet_pattern = r'^\s*[-–—•·→■▪□➤*]\s+'
    match = re.match(bullet_pattern, text)
    if match:
        return text[match.end():].strip()
    return text.strip()

def _remove_numbering(paragraph: Paragraph) -> None:
    """
    Полностью удаляет нумерацию из параграфа.

    Args:
        paragraph (Paragraph): Параграф, из которого нужно удалить нумерацию.
    """
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    
    num_pr = pPr.find(qn("w:numPr"))
    if num_pr is not None:
        pPr.remove(num_pr)
    
    for tag in [qn("w:ilvl"), qn("w:numId")]:
        elem = pPr.find(tag)
        if elem is not None:
            pPr.remove(elem)

def _set_outline_level(paragraph: Paragraph, level: int) -> None:
    """
    Устанавливает outline level для параграфа.

    Args:
        paragraph (Paragraph): Параграф, для которого нужно установить уровень.
        level (int): Уровень заголовка (1, 2, 3, 4).
    """
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    # Удаляем существующий outlineLvl, если он есть
    for existing_lvl in pPr.xpath('w:outlineLvl'):
        pPr.remove(existing_lvl)
    # Добавляем новый outlineLvl
    outlineLvl = OxmlElement('w:outlineLvl')
    outlineLvl.set(qn('w:val'), str(level - 1))  # outlineLvl: 0 для Heading 1, 1 для Heading 2 и т.д.
    pPr.append(outlineLvl)
    logger.debug(f"Установлен outline level {level - 1} для параграфа")

def _disable_auto_spacing(paragraph: Paragraph) -> None:
    """
    Отключает автоматические интервалы до и после абзаца (w:beforeAutospacing и w:afterAutospacing).

    Args:
        paragraph (Paragraph): Параграф, для которого нужно отключить автоинтервалы.
    """
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    
    # Отключаем автоматические интервалы
    spacing.set(qn('w:beforeAutospacing'), "0")
    spacing.set(qn('w:afterAutospacing'), "0")
    logger.debug("Отключены автоматические интервалы до и после абзаца")

def _clean_marker(text: str, marker_pattern: str) -> str:
    """
    Удаляет маркер из текста.

    Args:
        text (str): Текст для обработки.
        marker_pattern (str): Регулярное выражение для маркера.

    Returns:
        str: Текст без маркера.
    """
    return re.sub(marker_pattern, '', text).strip()

def format_document(
    elements: List[Dict[str, Any]],
    doc: Document,
    *,
    config: Optional[Dict[str, Any]] = None,
) -> Document:
    """
    Форматирует существующий документ Word, применяя стили к распознанным элементам,
    сохраняя при этом нераспознанные элементы (графики, диаграммы и т.д.).

    Args:
        elements: Список распознанных элементов документа.
        doc: Существующий документ Word.
        config: Настройки форматирования (replace_quotes, add_space_before, add_space_after, format_tables, add_space_after_heading1).

    Returns:
        Document: Отформатированный документ.
    """
    cfg = config or {}
    table_i = cfg.get("table_idx", 0)
    figure_i = cfg.get("figure_idx", 0)
    formula_i = cfg.get("formula_idx", 0)
    replace_quotes = cfg.get("replace_quotes", False)
    add_space_before = cfg.get("add_space_before", True)
    add_space_after = cfg.get("add_space_after", True)
    format_tables = cfg.get("format_tables", True)
    add_space_after_heading1 = cfg.get("add_space_after_heading1", True)
    table_row_height = cfg.get("table_row_height", 18)
    max_table_row_height = cfg.get("max_table_row_height", 60)

    cnt: Dict[int, int] = {}
    prev_list_type: str | None = None
    prev_level: int | None = None
    is_bibliography = False
    heading_counters = {1: 0, 2: 0, 3: 0, 4: 0}  # Сбрасываем счётчики заголовков

    element_dict = {el.get("para_idx"): el for el in elements if "para_idx" in el}

    para_idx = 0
    table_idx = 0
    body_elements = list(doc.element.body)
    i = 0
    while i < len(body_elements):
        node = body_elements[i]

        if node.tag == qn('w:p'):
            para = Paragraph(node, doc)

            if para_idx in element_dict:
                el = element_dict[para_idx]
                etype = el["type"]
                txt = el.get("text", "")

                if etype == "heading_1" and "список использованных источников" in txt.lower().strip():
                    is_bibliography = True
                    logger.debug("Начало раздела 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ'")
                    cnt = {}  # Сбрасываем счётчики для библиографии
                elif etype == "heading_1" and is_bibliography:
                    is_bibliography = False
                    logger.debug("Конец раздела 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ'")

                # ───── Заголовки ────────────────────────────────────
                if etype.startswith("heading"):
                    lvl = el.get("level", 1)
                    text = el.get("text", "")

                    text = _strip_heading_number(text, lvl)

                    if lvl == 1:
                        # Очищаем текст перед проверкой на NO_NUMERATION_HEADINGS
                        check_text = text.upper().strip()
                        if check_text in NO_NUMERATION_HEADINGS:
                            num = ""
                        else:
                            heading_counters[1] += 1
                            num = f"{heading_counters[1]}"  # Без точки в конце
                        heading_counters[2] = 0
                        heading_counters[3] = 0
                        heading_counters[4] = 0
                    elif lvl == 2:
                        heading_counters[2] += 1
                        heading_counters[3] = 0
                        heading_counters[4] = 0
                        num = f"{heading_counters[1]}.{heading_counters[2]}"  # Без точки в конце
                    elif lvl == 3:
                        heading_counters[3] += 1
                        heading_counters[4] = 0
                        num = f"{heading_counters[1]}.{heading_counters[2]}.{heading_counters[3]}"  # Без точки в конце
                    else:  # lvl == 4
                        heading_counters[4] += 1
                        num = f"{heading_counters[1]}.{heading_counters[2]}.{heading_counters[3]}.{heading_counters[4]}"  # Без точки в конце

                    if lvl == 1:
                        text = text.upper()

                    if lvl == 1:
                        break_para = para.insert_paragraph_before()
                        run = break_para.add_run()
                        run.add_break(WD_BREAK.PAGE)

                    # Создаём новый параграф с заголовком
                    new_para = para.insert_paragraph_before("", style=f"Heading {lvl}")
                    run = new_para.add_run(f"{num} {text}".strip())
                    
                    # Принудительно устанавливаем шрифт Times New Roman для всех run-объектов
                    for run in new_para.runs:
                        run.font.name = "Times New Roman"
                        logger.debug(f"Шрифт для заголовка уровня {lvl} установлен: Times New Roman")

                    # Устанавливаем outline level
                    _set_outline_level(new_para, lvl)

                    # Очищаем возможные конфликтующие настройки
                    _remove_numbering(new_para)

                    # Убедимся, что стиль применяется корректно
                    new_para.style = f"Heading {lvl}"

                    if lvl == 1 and add_space_after_heading1:
                        space_para = doc.add_paragraph("", style="Normal")
                        space_para.paragraph_format.line_spacing = 1.5  # Принудительно устанавливаем межстрочный интервал
                        space_para.paragraph_format.space_before = Pt(0)  # Принудительно устанавливаем отступ перед абзацем
                        space_para.paragraph_format.space_after = Pt(0)   # Принудительно устанавливаем отступ после абзаца
                        _disable_auto_spacing(space_para)  # Отключаем автоматические интервалы
                        new_para._element.addnext(space_para._element)

                    parent = para._element.getparent()
                    parent.remove(para._element)

                    prev_list_type = None
                    logger.debug(f"Добавлен новый заголовок уровня {lvl} на para_idx {para_idx}, старый удалён: '{text[:50]}'")

                # ───── Bullet-list ────────────────────────────────
                elif etype == "list_bullet":
                    lvl = el.get("list_level", 0)
                    core = el.get("text", "")
                    # Удаляем маркеры [UL], [UL:1]
                    core = _clean_marker(core, r'^\[(UL|UL:\d+)\]\s*')
                    core = _strip_bullet_marker(core)
                    # Очищаем параграф и добавляем новый текст
                    para.clear()
                    if replace_quotes:
                        core = _replace_quotes(core)
                    para.add_run(f"– {core}")
                    para.style = "List Bullet"
                    if lvl > 0:
                        para.paragraph_format.left_indent = Cm(1.25 * (lvl + 1))
                        para.paragraph_format.first_line_indent = Cm(-0.63)
                    prev_list_type, prev_level = "list_bullet", lvl
                    logger.debug(f"Отформатирован маркированный список на para_idx {para_idx}, уровень {lvl}: '{core[:50]}'")

                # ───── Ordered-list ────────────────────────────────
                elif etype == "list_ordered" and not is_bibliography:
                    lvl = el.get("list_level", 0)

                    if prev_list_type != "list_ordered" or (
                        prev_level is not None and lvl < prev_level
                    ):
                        cnt = {k: v for k, v in cnt.items() if k < lvl}

                    num = cnt.get(lvl, 0) + 1
                    cnt[lvl] = num

                    prefix = f"{'абвгдежзийклмнопрстуфхцчшщъыьэюя'[num-1]}) " if lvl == 0 else \
                             f"{num}) "
                    core = el.get("text", "")
                    # Удаляем маркеры [OL], [OL:1]
                    core = _clean_marker(core, r'^\[(OL|OL:\d+)\]\s*')
                    # Очищаем параграф и добавляем новый текст
                    para.clear()
                    if replace_quotes:
                        core = _replace_quotes(core)
                    para.add_run(f"{prefix}{core}")
                    para.style = "List Number"
                    if lvl > 0:
                        para.paragraph_format.left_indent = Cm(1.25 * (lvl + 1))
                        para.paragraph_format.first_line_indent = Cm(-0.63)
                    prev_list_type, prev_level = "list_ordered", lvl
                    logger.debug(f"Отформатирован нумерованный список на para_idx {para_idx}, уровень {lvl}: '{core[:50]}'")
                    
                # ───── Библиография ────────────────────────────────   
                elif (etype == "list_ordered" and is_bibliography) or is_bibliography:
                    lvl = el.get("list_level", 0) if not is_bibliography else 0

                    if prev_list_type != "list_ordered" or (
                        prev_level is not None and lvl < prev_level
                    ):
                        cnt = {k: v for k, v in cnt.items() if k < lvl}

                    num = cnt.get(lvl, 0) + 1
                    cnt[lvl] = num

                    prefix = f"{num}. " if lvl == 0 else \
                             f"{'абвгдежзийклмнопрстуфхцчшщъыьэюя'[num-1]}) "
                    core = el.get("text", "")
                    # Удаляем маркеры [OL], [OL:1]
                    core = _clean_marker(core, r'^\[(OL|OL:\d+)\]\s*')
                    # Очищаем параграф и добавляем новый текст
                    para.clear()
                    if replace_quotes:
                        core = _replace_quotes(core)
                    para.add_run(f"{prefix}{core}")
                    para.style = "List Number"
                    if lvl > 0:
                        para.paragraph_format.left_indent = Cm(1.25 * (lvl + 1))
                        para.paragraph_format.first_line_indent = Cm(-0.63)
                    prev_list_type, prev_level = "list_ordered", lvl
                    logger.debug(f"Отформатирован нумерованный список на para_idx {para_idx}, уровень {lvl}: '{core[:50]}'")

                # ───── Рисунок ─────────────────────────────────────
                elif etype == "figure":
                    if "number" not in el:
                        figure_i += 1
                        el["number"] = figure_i

                    img = el.get("image")
                    caption = el.get("caption", txt)

                    if add_space_before:
                        new_p = doc.add_paragraph("", style="Normal")
                        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        new_p.paragraph_format.line_spacing = 1.5  # Принудительно устанавливаем межстрочный интервал
                        new_p.paragraph_format.space_before = Pt(0)  # Принудительно устанавливаем отступ перед абзацем
                        new_p.paragraph_format.space_after = Pt(0)   # Принудительно устанавливаем отступ после абзаца
                        _disable_auto_spacing(new_p)  # Отключаем автоматические интервалы
                        node.addprevious(new_p._element)

                    p_img = doc.add_paragraph()
                    node.addprevious(p_img._element)

                    if img:
                        try:
                            buf = BytesIO()
                            img.save(buf, format="PNG")
                            buf.seek(0)
                            max_page_width = Cm(15.92)
                            dpi = 96
                            original_width_pixels = img.width
                            original_width_cm = Cm(original_width_pixels / dpi * 2.54)
                            if original_width_cm > max_page_width:
                                width = max_page_width
                            else:
                                width = original_width_cm
                            run = p_img.add_run()
                            run.add_picture(buf, width=width)
                            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_img.paragraph_format.left_indent = Cm(0)
                            p_img.paragraph_format.right_indent = Cm(0)
                            p_img.paragraph_format.first_line_indent = Cm(0)
                            p_img.paragraph_format.space_before = Pt(0)
                            p_img.paragraph_format.space_after = Pt(0)
                            logger.debug(f"Изображение №{el['number']} успешно добавлено в документ на para_idx {para_idx} с шириной {width.cm:.2f} см, отступы сброшены")
                        except Exception as e:
                            logger.error(f"Ошибка при добавлении изображения №{el['number']} на para_idx {para_idx}: {str(e)}")
                            caption = f"Ошибка: изображение №{el['number']} не добавлено"
                    else:
                        logger.warning(f"Изображение №{el['number']} отсутствует на para_idx {para_idx}, добавляется только подпись")

                    cap_txt = _norm_cap("figure", caption) if caption else f"Рисунок {el['number']}"
                    # Удаляем маркер [FIGURE]
                    cap_txt = _clean_marker(cap_txt, r'^\[FIGURE\]\s*')
                    if replace_quotes:
                        cap_txt = _replace_quotes(cap_txt)
                    p_cap = doc.add_paragraph(f"Рисунок {el['number']} — {cap_txt}", style="Caption")
                    p_img._element.addnext(p_cap._element)
                    logger.debug(f"Добавлена подпись рисунка №{el['number']} на para_idx {para_idx}: '{cap_txt}'")

                    if add_space_after:
                        new_p = doc.add_paragraph("", style="Normal")
                        new_p.paragraph_format.line_spacing = 1.5  # Принудительно устанавливаем межстрочный интервал
                        new_p.paragraph_format.space_before = Pt(0)  # Принудительно устанавливаем отступ перед абзацем
                        new_p.paragraph_format.space_after = Pt(0)   # Принудительно устанавливаем отступ после абзаца
                        _disable_auto_spacing(new_p)  # Отключаем автоматические интервалы
                        p_cap._element.addnext(new_p._element)

                    parent = para._element.getparent()
                    parent.remove(para._element)

                    if i + 1 < len(body_elements) and body_elements[i + 1].tag == qn('w:p'):
                        next_para = Paragraph(body_elements[i + 1], doc)
                        next_txt = next_para.text
                        if _is_caption_text(next_txt, "figure"):
                            next_parent = next_para._element.getparent()
                            next_parent.remove(next_para._element)
                            i += 1
                            para_idx += 1
                            logger.debug(f"Удалён исходный параграф подписи рисунка на para_idx {para_idx}: '{next_txt}'")
                        else:
                            para_idx += 1
                            i += 1
                    else:
                        para_idx += 1
                        i += 1

                    prev_list_type = None
                    continue

                # ───── Обычный текст ───────────────────────────────
                elif etype == "paragraph":
                    core = el.get("text", "")
                    if replace_quotes:
                        core = _replace_quotes(core)
                    para.clear()
                    para.add_run(core)
                    para.style = "Normal"
                    para.paragraph_format.line_spacing = 1.5  # Принудительно устанавливаем межстрочный интервал
                    para.paragraph_format.space_before = Pt(0)  # Принудительно устанавливаем отступ перед абзацем
                    para.paragraph_format.space_after = Pt(0)   # Принудительно устанавливаем отступ после абзаца
                    _disable_auto_spacing(para)  # Отключаем автоматические интервалы
                    prev_list_type = None
                    logger.debug(f"Отформатирован параграф на para_idx {para_idx}: '{core[:50]}'")

                # ───── Формула ─────────────────────────────────────
                elif etype == "formula":
                    if "number" not in el:
                        formula_i += 1
                        el["number"] = formula_i
                    txt = _clean_marker(txt, r'^\[FORMULA\]\s*')  # Удаляем маркер [FORMULA]
                    para.clear()
                    para.add_run(txt)
                    para.style = "Formula"
                    if f"({el['number']})" not in txt:
                        para.add_run(f" ({el['number']})")
                    prev_list_type = None
                    logger.debug(f"Отформатирована формула №{el['number']} на para_idx {para_idx}: '{txt[:50]}'")

            else:
                logger.debug(f"Элемент на para_idx {para_idx} не распознан, оставлен без изменений")

            para_idx += 1
            i += 1

        # ───── Таблица ─────────────────────────────────────
        elif node.tag == qn('w:tbl'):
            table_el = next((el for el in elements if el.get("type") == "table" and el.get("number") == table_i + 1), None)
            if table_el:
                table_i += 1

                if add_space_before:
                    new_p = doc.add_paragraph("", style="Normal")
                    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    new_p.paragraph_format.line_spacing = 1.5  # Принудительно устанавливаем межстрочный интервал
                    new_p.paragraph_format.space_before = Pt(0)  # Принудительно устанавливаем отступ перед абзацем
                    new_p.paragraph_format.space_after = Pt(0)   # Принудительно устанавливаем отступ после абзаца
                    _disable_auto_spacing(new_p)  # Отключаем автоматические интервалы
                    node.addprevious(new_p._element)

                cap_txt = _norm_cap("table", table_el.get("caption", ""))
                # Удаляем маркер [TABLE]
                cap_txt = _clean_marker(cap_txt, r'^\[TABLE\]\s*')
                if replace_quotes:
                    cap_txt = _replace_quotes(cap_txt)
                cap = doc.add_paragraph(f"Таблица {table_i} — {cap_txt}", style="Caption")
                cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
                node.addprevious(cap._element)
                logger.debug(f"Добавлена подпись таблицы №{table_i}: '{cap_txt}'")

                table = Table(node, doc)
                if format_tables:
                    table.style = "Table Grid"
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    table.autofit = False
                    for row in table.rows:
                        row.height = Pt(table_row_height)
                        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                    for c in table._cells:
                        for pr in c.paragraphs:
                            if replace_quotes:
                                full_text = "".join(run.text for run in pr.runs)
                                replaced_text = _replace_quotes(full_text)
                                pr.clear()
                                pr.add_run(replaced_text)
                            for run in pr.runs:
                                run.font.size = Pt(12)
                            pr.paragraph_format.first_line_indent = Cm(0)
                            pr.paragraph_format.left_indent = Cm(0)
                            pr.paragraph_format.space_before = Pt(0)
                            pr.paragraph_format.space_after = Pt(0)
                            pr.alignment = WD_ALIGN_PARAGRAPH.LEFT 
                    logger.debug(f"Отформатирована таблица №{table_i} с {len(table.rows)} строками")
                else:
                    logger.debug(f"Форматирование таблицы №{table_i} пропущено (format_tables=False)")

                if add_space_after:
                    new_p = doc.add_paragraph("", style="Normal")
                    new_p.paragraph_format.line_spacing = 1.5  # Принудительно устанавливаем межстрочный интервал
                    new_p.paragraph_format.space_before = Pt(0)  # Принудительно устанавливаем отступ перед абзацем
                    new_p.paragraph_format.space_after = Pt(0)   # Принудительно устанавливаем отступ после абзаца
                    _disable_auto_spacing(new_p)  # Отключаем автоматические интервалы
                    node.addnext(new_p._element)

            else:
                logger.debug(f"Таблица №{table_i + 1} не распознана, оставлена без изменений")

            table_idx += 1
            i += 1

        else:
            logger.debug(f"Элемент на позиции {i} (тип {node.tag}) не обрабатывается, оставлен без изменений")
            i += 1

    return doc